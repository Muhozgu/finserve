import json
import os
from datetime import datetime
from groq import Groq
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# CONFIG
GROQ_API_KEY = os.environ.get("GROQ_API_KEY", "gsk_xxxxx")  # Replace with your actual Groq API key or set as environment variable
GROQ_MODEL   = "llama-3.3-70b-versatile"
DATA_FILE    = "./portfolio_data.json"


# 1. Load data

def load_data(filepath: str) -> dict:
    with open(filepath, "r", encoding="utf-8") as f:
        return json.load(f)


# 2. Calculate key metrics

def calculate_metrics(data: dict) -> dict:
    cb   = data["source_core_banking"]
    crm  = data["source_crm"]
    prev = data["prior_month_comparison"]

    total     = cb["total_portfolio_value_pln"]
    overdue   = cb["total_overdue_pln"]
    written   = cb["total_written_off_pln"]
    prev_total   = prev["total_portfolio_value_pln"]
    prev_overdue = prev["total_overdue_pln"]

    par30 = (overdue / total * 100) if total else 0
    npl   = (written / total * 100) if total else 0
    collection_rate = (
        cb["total_repaid_this_month_pln"] /
        (cb["total_repaid_this_month_pln"] + overdue) * 100
    ) if (cb["total_repaid_this_month_pln"] + overdue) else 0

    portfolio_growth  = ((total - prev_total) / prev_total * 100) if prev_total else 0
    overdue_change    = ((overdue - prev_overdue) / prev_overdue * 100) if prev_overdue else 0
    loan_growth       = cb["total_active_loans"] - prev["total_active_loans"]
    disbursement_growth = (
        (cb["total_disbursed_this_month_pln"] - prev["total_disbursed_pln"]) /
        prev["total_disbursed_pln"] * 100
    ) if prev["total_disbursed_pln"] else 0

    return {
        "par30_pct":             round(par30, 2),
        "npl_pct":               round(npl, 2),
        "collection_rate_pct":   round(collection_rate, 2),
        "portfolio_growth_pct":  round(portfolio_growth, 2),
        "overdue_change_pct":    round(overdue_change, 2),
        "loan_growth":           loan_growth,
        "disbursement_growth_pct": round(disbursement_growth, 2),
    }


# 3. Build Groq prompt

SYSTEM_PROMPT = """You are the Chief Risk Officer at FinServe, a financial services
company specialising in SME lending. You write monthly portfolio reports for the
executive committee. Be analytical, direct, and concise.
Respond with ONLY a valid JSON object — no markdown, no code blocks, no extra text."""

def build_prompt(data: dict, metrics: dict) -> str:
    cb   = data["source_core_banking"]
    crm  = data["source_crm"]
    period = data["report_period"]

    return f"""
Analyse the following monthly portfolio data and respond with a JSON object
with exactly these keys:

{{
  "executive_summary": "3-4 sentence summary of portfolio performance this month",
  "key_findings": ["finding 1", "finding 2", "finding 3", "finding 4"],
  "risk_flags": ["risk 1", "risk 2", "risk 3"],
  "recommendations": ["recommendation 1", "recommendation 2", "recommendation 3"],
  "outlook": "2-3 sentence forward-looking commentary"
}}

PORTFOLIO DATA — {period['month']} {period['year']}:

Portfolio size:       PLN {cb['total_portfolio_value_pln']:,}  (prev: PLN {data['prior_month_comparison']['total_portfolio_value_pln']:,})
Active loans:         {cb['total_active_loans']}  (prev: {data['prior_month_comparison']['total_active_loans']})
New loans:            {cb['new_loans_this_month']}
Disbursed this month: PLN {cb['total_disbursed_this_month_pln']:,}
Repaid this month:    PLN {cb['total_repaid_this_month_pln']:,}
Total overdue:        PLN {cb['total_overdue_pln']:,}  (prev: PLN {data['prior_month_comparison']['total_overdue_pln']:,})
Written off:          PLN {cb['total_written_off_pln']:,}
Avg interest rate:    {cb['average_interest_rate_pct']}%

KEY RATIOS:
PAR30:                {metrics['par30_pct']}%
NPL ratio:            {metrics['npl_pct']}%
Collection rate:      {metrics['collection_rate_pct']}%
Portfolio growth MoM: {metrics['portfolio_growth_pct']}%
Overdue change MoM:   {metrics['overdue_change_pct']}%

OVERDUE BUCKETS:
1-30 days:   {cb['overdue_buckets'][0]['count']} loans  PLN {cb['overdue_buckets'][0]['value_pln']:,}
31-60 days:  {cb['overdue_buckets'][1]['count']} loans  PLN {cb['overdue_buckets'][1]['value_pln']:,}
61-90 days:  {cb['overdue_buckets'][2]['count']} loans  PLN {cb['overdue_buckets'][2]['value_pln']:,}
90+ days:    {cb['overdue_buckets'][3]['count']} loans

CRM DATA:
Active clients:         {crm['total_active_clients']}
New clients:            {crm['new_clients_this_month']}
Clients in restructuring: {crm['clients_in_restructuring']}
Clients flagged for review: {crm['clients_flagged_for_review']}
Pipeline value:         PLN {crm['pipeline_value_pln']:,}  ({crm['pipeline_deals']} deals)
"""

# 4. Call Groq
def call_groq(prompt: str) -> dict:
    client = Groq(api_key=GROQ_API_KEY)
    response = client.chat.completions.create(
        model=GROQ_MODEL,
        messages=[
            {"role": "system", "content": SYSTEM_PROMPT},
            {"role": "user",   "content": prompt}
        ],
        temperature=0.2
    )
    raw = response.choices[0].message.content.strip()
    return json.loads(raw)

# 5. Word document helpers
def add_divider(doc, color="2E74B5"):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"),   "single")
    bottom.set(qn("w:sz"),    "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color)
    pBdr.append(bottom)
    pPr.append(pBdr)
    p.paragraph_format.space_after = Pt(4)

def set_heading(paragraph, text, level=1):
    paragraph.clear()
    run = paragraph.add_run(text)
    run.bold = True
    if level == 1:
        run.font.size = Pt(13)
        run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
    else:
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)
    paragraph.paragraph_format.space_before = Pt(10)
    paragraph.paragraph_format.space_after  = Pt(4)

def shade_cell(cell, fill_hex):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  fill_hex)
    tcPr.append(shd)

def add_metric_table(doc, rows):
    table = doc.add_table(rows=len(rows), cols=3)
    table.style = "Table Grid"
    for i, (label, value, change) in enumerate(rows):
        table.cell(i, 0).text = label
        table.cell(i, 1).text = value
        table.cell(i, 2).text = change
        shade_cell(table.cell(i, 0), "DEEAF1")
        for j in range(3):
            for para in table.cell(i, j).paragraphs:
                for run in para.runs:
                    run.font.size = Pt(10)
    doc.add_paragraph()

def add_simple_table(doc, headers, rows, header_color="1F497D"):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    for j, h in enumerate(headers):
        cell = table.cell(0, j)
        cell.text = h
        shade_cell(cell, header_color)
        for para in cell.paragraphs:
            for run in para.runs:
                run.bold = True
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            cell = table.cell(i + 1, j)
            cell.text = str(val)
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(10)
            if i % 2 == 0:
                shade_cell(cell, "F5F9FC")
    doc.add_paragraph()

def add_bullet_list(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(item).font.size = Pt(10)

def add_flag_list(doc, items, color):
    for item in items:
        p = doc.add_paragraph()
        run = p.add_run(f"  {item}")
        run.font.size = Pt(10)
        run.font.color.rgb = color

# 6. Build Word document
def build_docx(data: dict, metrics: dict, commentary: dict, output_path: str):
    cb     = data["source_core_banking"]
    crm    = data["source_crm"]
    period = data["report_period"]
    prev   = data["prior_month_comparison"]
    doc    = Document()

    for section in doc.sections:
        section.top_margin    = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin   = Inches(1.2)
        section.right_margin  = Inches(1.2)

    # ── Title ────────────────────────────────────────────────────────────
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("MONTHLY PORTFOLIO REPORT")
    run.bold = True
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.add_run(
        f"{period['month']} {period['year']}  |  "
        f"Generated: {datetime.today().strftime('%d %B %Y')}"
    ).font.size = Pt(10)

    add_divider(doc)

    # ── 1. Executive Summary ─────────────────────────────────────────────
    h = doc.add_paragraph()
    set_heading(h, "1.  Executive Summary", level=1)
    p = doc.add_paragraph(commentary["executive_summary"])
    p.runs[0].font.size = Pt(10)

    # ── 2. Portfolio Overview ────────────────────────────────────────────
    h = doc.add_paragraph()
    set_heading(h, "2.  Portfolio Overview", level=1)

    add_metric_table(doc, [
        ("Total portfolio value",    f"PLN {cb['total_portfolio_value_pln']:,}",
         f"{metrics['portfolio_growth_pct']:+.1f}% MoM"),
        ("Active loans",             str(cb["total_active_loans"]),
         f"{metrics['loan_growth']:+d} vs prior month"),
        ("New loans this month",     str(cb["new_loans_this_month"]),
         f"prev: {prev['new_loans']}"),
        ("Disbursed this month",     f"PLN {cb['total_disbursed_this_month_pln']:,}",
         f"{metrics['disbursement_growth_pct']:+.1f}% MoM"),
        ("Repaid this month",        f"PLN {cb['total_repaid_this_month_pln']:,}", ""),
        ("Average loan size",        f"PLN {cb['average_loan_size_pln']:,}",       ""),
        ("Average interest rate",    f"{cb['average_interest_rate_pct']}%",        ""),
        ("Active clients",           str(crm["total_active_clients"]),
         f"+{crm['new_clients_this_month']} new this month"),
        ("Pipeline",                 f"PLN {crm['pipeline_value_pln']:,}",
         f"{crm['pipeline_deals']} deals"),
    ])

    # ── 3. Portfolio by Product ──────────────────────────────────────────
    h = doc.add_paragraph()
    set_heading(h, "3.  Portfolio by Product", level=1)

    add_simple_table(doc,
        headers=["Product", "No. of Loans", "Outstanding (PLN)"],
        rows=[
            (p["product"], p["count"], f"{p['value_pln']:,}")
            for p in cb["loan_book_by_product"]
        ]
    )

    # ── 4. Portfolio by Sector ───────────────────────────────────────────
    h = doc.add_paragraph()
    set_heading(h, "4.  Portfolio by Sector", level=1)

    add_simple_table(doc,
        headers=["Sector", "Clients", "Portfolio Value (PLN)"],
        rows=[
            (s["sector"], s["clients"], f"{s['portfolio_pln']:,}")
            for s in crm["sector_breakdown"]
        ]
    )

    # ── 5. Credit Quality & Risk Metrics ────────────────────────────────
    h = doc.add_paragraph()
    set_heading(h, "5.  Credit Quality & Risk Metrics", level=1)

    add_metric_table(doc, [
        ("PAR30 (Portfolio at Risk)",  f"{metrics['par30_pct']}%",
         f"overdue change: {metrics['overdue_change_pct']:+.1f}% MoM"),
        ("NPL ratio",                  f"{metrics['npl_pct']}%", ""),
        ("Collection rate",            f"{metrics['collection_rate_pct']}%", ""),
        ("Total overdue",              f"PLN {cb['total_overdue_pln']:,}",
         f"prev: PLN {prev['total_overdue_pln']:,}"),
        ("Written off this month",     f"PLN {cb['total_written_off_pln']:,}", ""),
        ("Clients in restructuring",   str(crm["clients_in_restructuring"]), ""),
        ("Clients flagged for review", str(crm["clients_flagged_for_review"]), ""),
    ])

    # Overdue buckets
    h2 = doc.add_paragraph()
    set_heading(h2, "Overdue buckets", level=2)

    add_simple_table(doc,
        headers=["Bucket", "No. of Loans", "Value (PLN)"],
        rows=[
            (b["bucket"], b["count"],
             f"{b['value_pln']:,}" if b["value_pln"] else "—")
            for b in cb["overdue_buckets"]
        ]
    )

    # ── 6. Top Exposures ─────────────────────────────────────────────────
    h = doc.add_paragraph()
    set_heading(h, "6.  Top Exposures", level=1)

    add_simple_table(doc,
        headers=["Client", "Loan ID", "Outstanding (PLN)", "Status"],
        rows=[
            (e["client"], e["loan_id"],
             f"{e['outstanding_pln']:,}", e["status"])
            for e in cb["top_exposures"]
        ]
    )

    # ── 7. Key Findings ──────────────────────────────────────────────────
    h = doc.add_paragraph()
    set_heading(h, "7.  Key Findings", level=1)
    add_bullet_list(doc, commentary["key_findings"])

    # ── 8. Risk Flags ────────────────────────────────────────────────────
    h = doc.add_paragraph()
    set_heading(h, "8.  Risk Flags", level=1)
    add_flag_list(doc, commentary["risk_flags"], RGBColor(0xC8, 0x23, 0x33))

    # ── 9. Recommendations ───────────────────────────────────────────────
    h = doc.add_paragraph()
    set_heading(h, "9.  Recommendations", level=1)
    add_bullet_list(doc, commentary["recommendations"])

    # ── 10. Outlook ──────────────────────────────────────────────────────
    h = doc.add_paragraph()
    set_heading(h, "10.  Outlook", level=1)
    p = doc.add_paragraph(commentary["outlook"])
    p.runs[0].font.size = Pt(10)

    # ── Footer ───────────────────────────────────────────────────────────
    add_divider(doc)
    footer = doc.add_paragraph(
        f"This report was automatically generated by FinServe AI Reporting Pipeline "
        f"on {datetime.today().strftime('%d %B %Y')}. "
        "For internal use only. Please review before distribution."
    )
    footer.runs[0].font.size = Pt(8)
    footer.runs[0].font.color.rgb = RGBColor(0x80, 0x80, 0x80)

    doc.save(output_path)
    print(f"\nReport saved to: {output_path}")

# 7. Main
def main():
    print("FinServe Monthly Portfolio Report Generator")
    print("============================================")

    print(f"Loading data from {DATA_FILE}...")
    data   = load_data(DATA_FILE)
    period = data["report_period"]
    print(f"Period: {period['month']} {period['year']}")

    print("Calculating metrics...")
    metrics = calculate_metrics(data)
    print(f"PAR30: {metrics['par30_pct']}%  |  Collection rate: {metrics['collection_rate_pct']}%  |  Portfolio growth: {metrics['portfolio_growth_pct']:+.1f}%")

    print("Calling Groq API for AI commentary...")
    prompt     = build_prompt(data, metrics)
    commentary = call_groq(prompt)
    print("AI commentary generated.")

    output_path = f"monthly_portfolio_report_{period['month']}_{period['year']}.docx"
    print("Building Word document...")
    build_docx(data, metrics, commentary, output_path)
    print("Done!")

if __name__ == "__main__":
    main()
