import json
import os
from datetime import datetime
from groq import Groq
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# CONFIG — paste your Groq API key here or set as environment variable
GROQ_API_KEY = os.environ.get("GROQ_API_KEY", "gsk_xxxxx")  # Replace with your actual Groq API key or set as environment variable
GROQ_MODEL   = "llama-3.3-70b-versatile"
DATA_FILE    = "client_data.json"

# 1. Load unified client data (simulates merged CRM + banking + loan output)
def load_client_data(filepath: str) -> dict:
    with open(filepath, "r", encoding="utf-8") as f:
        return json.load(f)

# 2. Build the prompt for Groq
SYSTEM_PROMPT = """You are a senior credit analyst at FinServe, a financial services 
company specialising in lending and credit products for SMEs.
Your job is to write clear, professional credit memos for the credit committee.
Be analytical, concise, and objective. Highlight both strengths and risks.
Always respond with ONLY a valid JSON object — no markdown, no code blocks, no extra text."""

def build_prompt(data: dict) -> str:
    crm   = data["source_crm"]
    bank  = data["source_core_banking"]
    loan  = data["source_loan_application"]
    fins  = bank["financials"]

    return f"""
Analyse the following client data and produce a credit memo as a JSON object
with exactly these keys:

{{
  "client_overview": "2-3 sentence summary of the company",
  "loan_request": "2-3 sentence description of what is being requested and why",
  "financial_summary": "3-4 sentences covering revenue trend, EBITDA, profitability and leverage",
  "key_ratios": {{
    "debt_to_ebitda_2023": <number>,
    "ebitda_margin_2023": <number as percentage>,
    "revenue_growth_yoy": <number as percentage>,
    "debt_to_equity_2023": <number>
  }},
  "credit_strengths": ["strength 1", "strength 2", "strength 3"],
  "credit_risks": ["risk 1", "risk 2", "risk 3"],
  "collateral_assessment": "2 sentence assessment of the collateral offered",
  "recommendation": "APPROVE or DECLINE",
  "recommendation_rationale": "2-3 sentence rationale for the recommendation",
  "conditions": ["condition 1", "condition 2"]
}}

CLIENT DATA:

--- CRM ---
Company:         {crm['company_name']}
Legal form:      {crm['legal_form']}
Incorporated:    {crm['incorporation_date']}
Industry:        {crm['industry']}
Contact:         {crm['contact_name']}, {crm['contact_title']}
Account manager: {crm['account_manager']}
Notes:           {crm['relationship_notes']}

--- CORE BANKING ---
Credit score:    {bank['credit_bureau_score']}
Payment history: {bank['payment_history']}
Existing debt:   PLN {bank['existing_facilities'][0]['outstanding_pln']:,} ({bank['existing_facilities'][0]['lender']}, {bank['existing_facilities'][0]['status']})

Financials (PLN):
Year       Revenue      EBITDA       Net Profit   Total Debt   Equity
2023       {fins['year_2023']['revenue']:>12,} {fins['year_2023']['ebitda']:>12,} {fins['year_2023']['net_profit']:>12,} {fins['year_2023']['total_debt']:>12,} {fins['year_2023']['equity']:>12,}
2022       {fins['year_2022']['revenue']:>12,} {fins['year_2022']['ebitda']:>12,} {fins['year_2022']['net_profit']:>12,} {fins['year_2022']['total_debt']:>12,} {fins['year_2022']['equity']:>12,}
2021       {fins['year_2021']['revenue']:>12,} {fins['year_2021']['ebitda']:>12,} {fins['year_2021']['net_profit']:>12,} {fins['year_2021']['total_debt']:>12,} {fins['year_2021']['equity']:>12,}

--- LOAN APPLICATION ---
Application ID:  {loan['application_id']}
Product:         {loan['product_type']}
Amount:          PLN {loan['amount_pln']:,}
Term:            {loan['requested_term_months']} months
Purpose:         {loan['purpose']}
Collateral:      {loan['collateral_offered']}
Documents:       {', '.join(loan['supporting_documents'])}
"""

# 3. Call Groq API
def call_groq(prompt: str) -> dict:
    client = Groq(api_key=GROQ_API_KEY)
    response = client.chat.completions.create(
        model=GROQ_MODEL,
        messages=[
            {"role": "system", "content": SYSTEM_PROMPT},
            {"role": "user",   "content": prompt}
        ],
        temperature=0.2
    )
    raw = response.choices[0].message.content.strip()
    return json.loads(raw)

# 4. Word document helpers
def set_heading_style(paragraph, text, level=1):
    paragraph.clear()
    run = paragraph.add_run(text)
    run.bold = True
    if level == 1:
        run.font.size = Pt(14)
        run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
    else:
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)
    paragraph.paragraph_format.space_before = Pt(12)
    paragraph.paragraph_format.space_after  = Pt(4)

def add_divider(doc):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "2E74B5")
    pBdr.append(bottom)
    pPr.append(pBdr)
    p.paragraph_format.space_after = Pt(6)

def add_kv_table(doc, rows: list):
    table = doc.add_table(rows=len(rows), cols=2)
    table.style = "Table Grid"
    for i, (key, val) in enumerate(rows):
        table.cell(i, 0).text = key
        table.cell(i, 1).text = str(val)
        for cell in table.row_cells(i):
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(10)
        # shade label column
        tc = table.cell(i, 0)._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"),   "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"),  "DEEAF1")
        tcPr.append(shd)
    doc.add_paragraph()

def add_bullet_list(doc, items: list):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(item).font.size = Pt(10)

# 5. Build the Word document
def build_docx(data: dict, memo: dict, output_path: str):
    crm  = data["source_crm"]
    loan = data["source_loan_application"]
    fins = data["source_core_banking"]["financials"]
    doc  = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin    = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin   = Inches(1.2)
        section.right_margin  = Inches(1.2)

    # ── Title block ──────────────────────────────────────────────────────
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("CREDIT MEMORANDUM")
    run.bold = True
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.add_run(
        f"{crm['company_name']}  |  {loan['application_id']}  |  "
        f"Prepared: {datetime.today().strftime('%d %B %Y')}"
    ).font.size = Pt(10)

    add_divider(doc)

    # ── 1. Application summary table ────────────────────────────────────
    h = doc.add_paragraph()
    set_heading_style(h, "1.  Application Summary", level=1)

    add_kv_table(doc, [
        ("Company",          crm["company_name"]),
        ("Legal form",       crm["legal_form"]),
        ("Industry",         crm["industry"]),
        ("Incorporated",     crm["incorporation_date"]),
        ("Account manager",  crm["account_manager"]),
        ("Product",          loan["product_type"]),
        ("Amount requested", f"PLN {loan['amount_pln']:,}"),
        ("Term",             f"{loan['requested_term_months']} months"),
        ("Repayment",        loan["repayment_type"]),
        ("Purpose",          loan["purpose"]),
    ])

    # ── 2. Client overview ───────────────────────────────────────────────
    h = doc.add_paragraph()
    set_heading_style(h, "2.  Client Overview", level=1)
    p = doc.add_paragraph(memo["client_overview"])
    p.runs[0].font.size = Pt(10)

    # ── 3. Loan request ──────────────────────────────────────────────────
    h = doc.add_paragraph()
    set_heading_style(h, "3.  Loan Request", level=1)
    p = doc.add_paragraph(memo["loan_request"])
    p.runs[0].font.size = Pt(10)

    # ── 4. Financial analysis ────────────────────────────────────────────
    h = doc.add_paragraph()
    set_heading_style(h, "4.  Financial Analysis", level=1)

    h2 = doc.add_paragraph()
    set_heading_style(h2, "Historical financials (PLN)", level=2)

    fin_table = doc.add_table(rows=5, cols=4)
    fin_table.style = "Table Grid"
    headers = ["", "2021", "2022", "2023"]
    rows_data = [
        ("Revenue",    fins["year_2021"]["revenue"],    fins["year_2022"]["revenue"],    fins["year_2023"]["revenue"]),
        ("EBITDA",     fins["year_2021"]["ebitda"],     fins["year_2022"]["ebitda"],     fins["year_2023"]["ebitda"]),
        ("Net profit", fins["year_2021"]["net_profit"], fins["year_2022"]["net_profit"], fins["year_2023"]["net_profit"]),
        ("Total debt", fins["year_2021"]["total_debt"], fins["year_2022"]["total_debt"], fins["year_2023"]["total_debt"]),
    ]

    for j, h_text in enumerate(headers):
        cell = fin_table.cell(0, j)
        cell.text = h_text
        for para in cell.paragraphs:
            for run in para.runs:
                run.bold = True
                run.font.size = Pt(10)
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"),   "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"),  "1F497D")
        tcPr.append(shd)
        for para in cell.paragraphs:
            for run in para.runs:
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    for i, (label, v21, v22, v23) in enumerate(rows_data):
        fin_table.cell(i+1, 0).text = label
        fin_table.cell(i+1, 1).text = f"{v21:,}"
        fin_table.cell(i+1, 2).text = f"{v22:,}"
        fin_table.cell(i+1, 3).text = f"{v23:,}"
        for j in range(4):
            for para in fin_table.cell(i+1, j).paragraphs:
                for run in para.runs:
                    run.font.size = Pt(10)

    doc.add_paragraph()

    h2 = doc.add_paragraph()
    set_heading_style(h2, "Key ratios (2023)", level=2)
    ratios = memo["key_ratios"]
    add_kv_table(doc, [
        ("Debt / EBITDA",      f"{ratios['debt_to_ebitda_2023']:.2f}x"),
        ("EBITDA margin",      f"{ratios['ebitda_margin_2023']:.1f}%"),
        ("Revenue growth YoY", f"{ratios['revenue_growth_yoy']:.1f}%"),
        ("Debt / Equity",      f"{ratios['debt_to_equity_2023']:.2f}x"),
    ])

    h2 = doc.add_paragraph()
    set_heading_style(h2, "Financial commentary", level=2)
    p = doc.add_paragraph(memo["financial_summary"])
    p.runs[0].font.size = Pt(10)

    # ── 5. Credit assessment ─────────────────────────────────────────────
    h = doc.add_paragraph()
    set_heading_style(h, "5.  Credit Assessment", level=1)

    h2 = doc.add_paragraph()
    set_heading_style(h2, "Strengths", level=2)
    add_bullet_list(doc, memo["credit_strengths"])

    h2 = doc.add_paragraph()
    set_heading_style(h2, "Risks", level=2)
    add_bullet_list(doc, memo["credit_risks"])

    # ── 6. Collateral ────────────────────────────────────────────────────
    h = doc.add_paragraph()
    set_heading_style(h, "6.  Collateral", level=1)
    add_kv_table(doc, [("Collateral offered", loan["collateral_offered"])])
    p = doc.add_paragraph(memo["collateral_assessment"])
    p.runs[0].font.size = Pt(10)

    # ── 7. Recommendation ────────────────────────────────────────────────
    add_divider(doc)
    h = doc.add_paragraph()
    set_heading_style(h, "7.  Recommendation", level=1)

    rec_para = doc.add_paragraph()
    rec_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rec_run = rec_para.add_run(f"  {memo['recommendation']}  ")
    rec_run.bold = True
    rec_run.font.size = Pt(16)
    color = RGBColor(0x1E, 0x7E, 0x34) if memo["recommendation"] == "APPROVE" else RGBColor(0xC8, 0x23, 0x33)
    rec_run.font.color.rgb = color

    p = doc.add_paragraph(memo["recommendation_rationale"])
    p.runs[0].font.size = Pt(10)

    if memo.get("conditions"):
        h2 = doc.add_paragraph()
        set_heading_style(h2, "Conditions", level=2)
        add_bullet_list(doc, memo["conditions"])

    # ── Footer note ──────────────────────────────────────────────────────
    add_divider(doc)
    footer_p = doc.add_paragraph(
        f"This memorandum was prepared by {crm['account_manager']} on "
        f"{datetime.today().strftime('%d %B %Y')} using FinServe AI Credit Memo Generator. "
        "It is intended for internal use by the credit committee only."
    )
    footer_p.runs[0].font.size = Pt(8)
    footer_p.runs[0].font.color.rgb = RGBColor(0x80, 0x80, 0x80)

    doc.save(output_path)
    print(f"\nCredit memo saved to: {output_path}")

# 6. Main
def main():
    print("FinServe Credit Memo Generator")
    print("================================")

    print(f"Loading client data from {DATA_FILE}...")
    data = load_client_data(DATA_FILE)
    company = data["source_crm"]["company_name"]
    app_id  = data["source_loan_application"]["application_id"]
    print(f"Client: {company} | Application: {app_id}")

    print("Calling Groq API (llama3-70b)...")
    prompt = build_prompt(data)
    memo   = call_groq(prompt)
    print("AI analysis complete.")
    print(f"Recommendation: {memo['recommendation']}")

    safe_name   = company.replace(" ", "_").replace(".", "").replace("/", "")
    output_path = f"credit_memo_{safe_name}.docx"

    print("Building Word document...")
    build_docx(data, memo, output_path)
    print("Done!")

if __name__ == "__main__":
    main()
